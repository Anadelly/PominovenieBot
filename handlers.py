import os
import re
import logging
from datetime import datetime
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ZAPISKI_DIR = "zapiski"
ZDRAVIE_FILE = os.path.join(ZAPISKI_DIR, "о_здравии.docx")
UPOKOENIE_FILE = os.path.join(ZAPISKI_DIR, "о_упокоении.docx")

def ensure_dir():
    os.makedirs(ZAPISKI_DIR, exist_ok=True)

def append_to_docx(filepath, names, sender):
    from docx import Document

    if os.path.exists(filepath):
        doc = Document(filepath)
    else:
        doc = Document()
        # Заголовок
        heading = "О здравии" if "здрав" in filepath else "Об упокоении"
        p = doc.add_paragraph(heading)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].font.size = Pt(12)

    # Добавим имя пользователя
    doc.add_paragraph(f"{sender}:", style=None).runs[0].font.size = Pt(12)

    # Формируем абзац с тремя колонками
    column_count = 4
    columns = [[] for _ in range(column_count)]

    for i, name in enumerate(names):
        columns[i % column_count].append(name)

    max_len = max(len(col) for col in columns)
    for i in range(max_len):
        row = []
        for col in columns:
            row.append(col[i] if i < len(col) else "")
        line = "     ".join(row)  # пробелы между колонками
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(12)

    doc.save(filepath)

def get_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("О здравии", callback_data="ozdravii")],
        [InlineKeyboardButton("Об упокоении", callback_data="oupokoenii")],
        [InlineKeyboardButton("Пожертвовать", callback_data="donate")],
        [InlineKeyboardButton("🔁 Начать заново", callback_data="restart")]
    ])

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("О здравии", callback_data="ozdravii")],
        [InlineKeyboardButton("Об упокоении", callback_data="oupokoenii")],
        [InlineKeyboardButton("Пожертвовать", callback_data="donate")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        "👋 Добро пожаловать! Этот бот помогает подать записки в православный храм свщмч. Сильвестра Омского г. Омска 📜\n"
        "\n"
        "Вы можете выбрать тип записки и ввести имена. Также можно внести пожертвование на своё усмотрение.\n"
        "Имена будут помянуты на ближайшей Литургии 🙏\n"
        "\n"
        "Выберите действие:",
        reply_markup=reply_markup
    )


async def button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if query.data in ["ozdravii", "oupokoenii"]:
        context.user_data["type"] = query.data
        await query.message.reply_text("Введите только имена в родительном падеже через пробел, без дополнительных знаков и символов: Марии Сергия Анатолия")
    elif query.data == "donate":
        with open("static/qr-code.jpg", "rb") as qr:
            await query.message.reply_photo(
                photo=qr,
                caption="Введите сумму в приложении банка по ссылке https://qr.nspk.ru/BS2A003TTV82T23F844A34OJIMUM20JS?type=01&bank=100000000005&crc=7FF6 Спасибо!"
            )
        await query.message.reply_text("Выберите действие:", reply_markup=get_keyboard())
    elif query.data == "restart":
        context.user_data.clear()
        await query.edit_message_text(
            "🔄 Начнём заново! Пожалуйста, выберите действие:",
            reply_markup=get_keyboard()
        )    

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    note_type = context.user_data.get("type")
    if not note_type:
        return

    ensure_dir()
    names = re.findall(r"[А-Яа-яёЁ ]+", update.message.text)
    names = [name.strip() for name in names if name.strip()]
    sender = update.effective_user.full_name or "неизвестный"

    if names:
        filepath = ZDRAVIE_FILE if note_type == "ozdravii" else UPOKOENIE_FILE
        append_to_docx(filepath, names, sender)
        response = "✅ Имена о здравии приняты." if note_type == "ozdravii" else "✅ Имена об упокоении приняты."
        await update.message.reply_text(response, reply_markup=get_keyboard())
    else:
        await update.message.reply_text("⚠ Не удалось распознать имена. Попробуйте снова.", reply_markup=get_keyboard())
    context.user_data["type"] = None

async def export_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    from telegram.constants import ChatAction
    import yadisk

    user_id = update.effective_user.id
    admin_ids = [int(i.strip()) for i in os.getenv("ADMIN_IDS", "").split(",") if i.strip()]
    if user_id not in admin_ids:
        await update.message.reply_text("⛔ У вас нет доступа к этой команде.")
        return

    ensure_dir()
    exported = False
    for path in [ZDRAVIE_FILE, UPOKOENIE_FILE]:
        if os.path.exists(path):
            await update.message.chat.send_action(action=ChatAction.UPLOAD_DOCUMENT)
            with open(path, "rb") as f:
                await update.message.reply_document(f)
            try:
                token = os.getenv("YANDEX_DISK_TOKEN")
                if token:
                    y = yadisk.YaDisk(token=token)
                    y.upload(path, f"/pominovenie/{os.path.basename(path)}", overwrite=True)
            except Exception as e:
                logging.error(f"Ошибка при выгрузке на Яндекс.Диск: {e}")
            os.remove(path)
            exported = True

    if exported:
        await update.message.reply_text("✅ Записки выгружены, скопированы на Яндекс.Диск и обнулены.")
    else:
        await update.message.reply_text("⚠ Папка с записками пуста.")


from utils.yadisk_utils import upload_docx_to_yadisk
from datetime import datetime
import os
import shutil

async def upload_yadisk_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    ADMIN_IDS = list(map(int, os.environ.get("ADMIN_IDS", "").split(",")))
    if update.effective_user.id not in ADMIN_IDS:
        return
    filename = f"zapiski_{datetime.now().strftime('%Y-%m-%d')}.docx"
    file_path = os.path.join(".", filename)
    if not os.path.exists(file_path):
        await update.message.reply_text("Файл не найден для загрузки.")
        return
    archived_path = os.path.join(".", f"арх_{filename}")
    shutil.move(file_path, archived_path)
    success = upload_docx_to_yadisk(archived_path)
    if success:
        await update.message.reply_text("Файл успешно загружен на Яндекс.Диск.")
    else:
        await update.message.reply_text("Ошибка при загрузке на Яндекс.Диск.")
