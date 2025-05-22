import os
import re
from datetime import datetime
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ContextTypes,
    filters
)
from docx import Document
from docx.shared import Pt

# Конфигурация из переменных окружения
BOT_TOKEN = os.getenv("BOT_TOKEN")
ADMIN_IDS = [int(i.strip()) for i in os.getenv("ADMIN_IDS", "").split(",") if i.strip()]

# === Команды и обработчики ===

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("О здравии", callback_data='ozdravii')],
        [InlineKeyboardButton("Об упокоении", callback_data='oupokoenii')],
        [InlineKeyboardButton("Пожертвовать", callback_data='donate')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Выберите действие:", reply_markup=reply_markup)

async def button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    if query.data in ['ozdravii', 'oupokoenii']:
        context.user_data['type'] = query.data
        await query.message.reply_text("Пожалуйста, введите имена в формате:\n\nболящей Марии, младенца Сергия, воина Александра")
    
    elif query.data == 'donate':
        with open('static/qr-code.jpg', 'rb') as qr:
            await query.message.reply_photo(
                photo=qr,
                caption="📷 Отсканируйте QR-код в приложении банка и введите сумму перевода.\n\n💖 Спасибо за помощь!"
            )

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    note_type = context.user_data.get('type')
    if note_type:
        names = re.findall(r'[А-Яа-яёЁ ]+', update.message.text)
        names = [name.strip() for name in names if name.strip()]

        if names:
            now = datetime.now().strftime("%d%m%Y")
            filename = f"{'о_здравии' if note_type == 'ozdravii' else 'о_упокоении'}_{now}.docx"
            filepath = os.path.join("zapiski", filename)

            os.makedirs("zapiski", exist_ok=True)
            doc = Document()
            table = doc.add_table(rows=1, cols=3)
            cells = table.rows[0].cells

            for i in range(min(3, len(names))):
                paragraph = cells[i].paragraphs[0]
                run = paragraph.add_run(names[i])
                run.font.size = Pt(14)

            doc.save(filepath)
            await update.message.reply_text("Записка сохранена. 🙏")
        else:
            await update.message.reply_text("Не удалось распознать имена. Попробуйте снова.")
        
        context.user_data['type'] = None

# === Запуск бота ===

def main():
    application = ApplicationBuilder().token(BOT_TOKEN).build()
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CallbackQueryHandler(button))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    application.run_polling()

if __name__ == '__main__':
    main()
