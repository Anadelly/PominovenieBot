
import json
from docx import Document
import os

DATA_FILE = "notes_data.json"

def save_note(note_type, names, username):
    if not os.path.exists(DATA_FILE):
        data = {"health": [], "repose": []}
    else:
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)

    data[note_type].append({
        "names": names,
        "user": username
    })

    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_notes():
    if not os.path.exists(DATA_FILE):
        return {"health": [], "repose": []}
    with open(DATA_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def clear_notes():
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump({"health": [], "repose": []}, f)

def generate_docx(file_path):
    data = load_notes()
    doc = Document()
    doc.add_heading("О здравии", level=1)
    for entry in data["health"]:
        doc.add_paragraph(f"{entry['names']} (от {entry['user']})")
    doc.add_heading("Об упокоении", level=1)
    for entry in data["repose"]:
        doc.add_paragraph(f"{entry['names']} (от {entry['user']})")
    doc.save(file_path)
